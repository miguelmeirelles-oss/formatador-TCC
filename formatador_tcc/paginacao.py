"""Numeração de página -- Apêndice I, seção PAGINAÇÃO:

"Todas as folhas do trabalho, a partir da folha de rosto, devem ser
contadas sequencialmente, mas não numeradas. A numeração deve ser colocada
a partir da primeira folha da parte textual (Introdução), em algarismos
arábicos, no canto superior direito da folha. Havendo anexo(s) e
apêndice(s), as suas folhas devem ser numeradas e paginadas de maneira
contínua."

O comentário da própria autora do modelo oficial detalha o cálculo: contam-se
todas as páginas físicas anteriores exceto a CAPA e a FICHA CATALOGRÁFICA --
por isso o número mostrado na Introdução é sempre "página física atual menos
2" (config.PAGINAS_EXCLUIDAS_DA_CONTAGEM).

Assim como o Sumário, não há como calcular a paginação real sem um motor de
layout de verdade (python-docx não faz isso). A solução é a mesma: um campo
nativo do Word -- aqui um campo de fórmula `{ =PAGE-2 }` --, que o próprio
Word recalcula automaticamente ao abrir o documento (settings.xml já pede
atualização de campos ao abrir, configurado por sumario.py).
"""
from __future__ import annotations

import copy
from dataclasses import dataclass

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .classify import EstadoClassificacao, classificar_paragrafo
from .config import PAGINAS_EXCLUIDAS_DA_CONTAGEM


@dataclass
class ResultadoPaginacao:
    aplicada: bool
    secao_introducao: int | None = None


def _mapear_paragrafo_para_secao(document) -> list[int]:
    """Para cada parágrafo de nível superior (mesma ordem de
    `document.paragraphs`), a que seção (`document.sections`) ele pertence."""
    body = document.element.body
    mapa: list[int] = []
    secao_atual = 0
    for child in body:
        if child.tag == qn("w:p"):
            mapa.append(secao_atual)
            ppr = child.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                secao_atual += 1
    return mapa


def _remover_reinicio_numeracao(sectPr) -> None:
    """Remove `w:pgNumType/@w:start` (reinício de contagem de página) de uma
    seção. A contagem de página do documento deve ser contínua do começo ao
    fim (ver docstring do módulo) -- se qualquer seção reiniciar em 1, o
    campo `{ =PAGE-2 }` inserido na Introdução passa a contar a partir
    dessa seção, não do documento inteiro."""
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        return
    if pgNumType.get(qn("w:start")) is not None:
        del pgNumType.attrib[qn("w:start")]
    if len(pgNumType.attrib) == 0:
        sectPr.remove(pgNumType)


def _garantir_secao_propria_a_partir_de(document, indice_paragrafo: int) -> None:
    """Garante que o parágrafo em `indice_paragrafo` seja o primeiro de uma
    seção própria, inserindo uma quebra de seção (nova página) logo antes
    dele se ainda não houver uma.

    Muitos alunos digitam o trabalho inteiro numa única seção do Word (sem
    preservar as quebras de seção do modelo oficial) -- sem uma seção
    própria a partir da Introdução, não há como o cabeçalho (onde fica o
    número de página) ser diferente antes e depois dela, porque cabeçalho é
    uma propriedade de seção, não de página. Isso também garante que a
    Introdução comece numa página nova, conforme o Apêndice I.
    """
    if indice_paragrafo <= 0:
        return

    mapa = _mapear_paragrafo_para_secao(document)
    if indice_paragrafo >= len(mapa):
        return
    if mapa[indice_paragrafo] != mapa[indice_paragrafo - 1]:
        return  # já é o início de uma seção

    paragrafos = document.paragraphs
    anterior = paragrafos[indice_paragrafo - 1]
    secao_atual = document.sections[mapa[indice_paragrafo]]

    nova_sectPr = copy.deepcopy(secao_atual._sectPr)
    # a cópia não pode manter as referências de cabeçalho/rodapé da seção
    # original -- senão as duas seções passam a apontar para a mesma parte
    # do pacote (mesmo r:id), e mexer no cabeçalho de uma mexe na outra
    # também. A seção nova (a que fica ANTES da quebra) começa "vinculada
    # à anterior" (sem cabeçalho próprio), que é o padrão do Word.
    for tag in (qn("w:headerReference"), qn("w:footerReference")):
        for ref in nova_sectPr.findall(tag):
            nova_sectPr.remove(ref)

    ppr = anterior._p.get_or_add_pPr()
    ppr.append(nova_sectPr)

    # a numeração de página (ver aplicar_numeracao_paginas) depende da
    # contagem de página FÍSICA e contínua do documento inteiro -- se esta
    # seção nova (ou a original, copiada logo acima) reiniciar a contagem em
    # 1, o campo PAGE da seção da Introdução passa a contar a partir do
    # início dela mesma, não do documento inteiro, e "=PAGE-2" mostra um
    # número errado (às vezes até negativo).
    _remover_reinicio_numeracao(nova_sectPr)


def _indice_primeiro_titulo1(document) -> int | None:
    estado = EstadoClassificacao()
    for i, paragraph in enumerate(document.paragraphs):
        categoria = classificar_paragrafo(paragraph, estado)
        if categoria == "titulo1":
            return i
    return None


def _construir_campo_numero_pagina(paragraph) -> None:
    """Substitui o conteúdo do parágrafo por um campo `{ = { PAGE } - N }`.

    Um campo de fórmula (`{ = ... }`) só reconhece PAGE como o campo nativo
    de número de página se ele for ANINHADO como um campo de verdade (outro
    par begin/instrText/end dentro da instrução da fórmula) -- escrever
    "PAGE" como texto simples dentro de `{ =PAGE-2 }` faz o Word tratar
    "PAGE" como o nome de um indicador (bookmark) inexistente, e o campo
    mostra "Erro! Indicador não definido." em vez do número.
    """
    p = paragraph._p
    for child in list(p.findall(qn("w:r"))):
        p.remove(child)

    def novo_run():
        r = OxmlElement("w:r")
        p.append(r)
        return r

    def instrText(texto):
        r = novo_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = texto
        r.append(instr)

    def fldChar(tipo):
        r = novo_run()
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), tipo)
        if tipo == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)

    fldChar("begin")           # campo externo (fórmula) -- begin
    instrText(" = ")
    fldChar("begin")           # campo interno (PAGE) -- begin
    instrText(" PAGE ")
    fldChar("end")             # campo interno (PAGE) -- end
    instrText(f" - {PAGINAS_EXCLUIDAS_DA_CONTAGEM} ")
    fldChar("separate")        # campo externo -- separator (resultado em cache abaixo)

    r_resultado = novo_run()
    t = OxmlElement("w:t")
    t.text = "1"
    r_resultado.append(t)

    fldChar("end")              # campo externo -- end


def aplicar_numeracao_paginas(document) -> ResultadoPaginacao:
    indice_intro = _indice_primeiro_titulo1(document)
    if indice_intro is None:
        return ResultadoPaginacao(aplicada=False)

    _garantir_secao_propria_a_partir_de(document, indice_intro)

    # o documento do aluno pode já trazer, na sua única seção original, um
    # reinício de numeração (comum: o Word grava `w:start="1"` por padrão
    # mesmo quando o autor nunca mexeu nisso) -- precisa ser removido de
    # TODAS as seções, não só da que acabou de ser criada acima, senão a
    # seção da Introdução conta a partir de si mesma em vez do documento
    # inteiro.
    for s in document.sections:
        _remover_reinicio_numeracao(s._sectPr)

    mapa_secoes = _mapear_paragrafo_para_secao(document)
    if indice_intro >= len(mapa_secoes):
        return ResultadoPaginacao(aplicada=False)
    secao_idx = mapa_secoes[indice_intro]

    secoes = document.sections
    if secao_idx >= len(secoes):
        return ResultadoPaginacao(aplicada=False)

    # Garante que nenhuma seção pré-textual mostre número de página, mesmo
    # que o documento já tivesse algo (errado) configurado ali antes --
    # por exemplo, reprocessar um arquivo que passou por uma versão antiga
    # desta ferramenta e ficou com um campo de página na seção errada. Só
    # esvazia o conteúdo (cabeçalho em branco = nada aparece na página,
    # independente do que a seção 0 reporta para is_linked_to_previous --
    # esse flag tem um comportamento instável especificamente na primeira
    # seção do documento).
    for s in secoes[:secao_idx]:
        for child in list(s.header._element):
            s.header._element.remove(child)

    secao = secoes[secao_idx]
    secao.header.is_linked_to_previous = False
    header = secao.header

    # remove qualquer conteúdo pré-existente no cabeçalho -- inclusive
    # blocos de "Número de página" prontos do Word (`w:sdt`), que não
    # aparecem em header.paragraphs mas ficariam duplicando o número.
    for child in list(header._element):
        header._element.remove(child)
    paragrafo = header.add_paragraph()

    paragrafo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _construir_campo_numero_pagina(paragrafo)

    # Seções seguintes (Referências, Apêndices, Anexos) herdam este mesmo
    # cabeçalho -- numeração contínua, conforme o Apêndice I.
    for s in secoes[secao_idx + 1:]:
        s.header.is_linked_to_previous = True

    return ResultadoPaginacao(aplicada=True, secao_introducao=secao_idx)
