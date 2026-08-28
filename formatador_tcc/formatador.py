"""Motor de normalização de formatação ABNT.

`formatar_documento` recebe um documento python-docx já aberto (o TCC do
aluno) e aplica a formatação ABNT do modelo oficial em cada parágrafo do
corpo do documento -- fonte, tamanho, alinhamento, espaçamento, recuo -- sem
jamais reescrever o texto dos parágrafos. Retorna uma lista de eventos
(auditoria do que foi tocado) para entrar no relatório final.
"""
from __future__ import annotations

from dataclasses import dataclass, replace
from typing import Optional

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from . import config
from .classify import (
    EstadoClassificacao,
    classificar_paragrafo,
    eh_marcador_ficha_catalografica,
    eh_nome_de_estilo_titulo_numerado,
    em_zona_pre_textual,
)
from .estilos import aplicar_configuracao_pagina, aplicar_estilo, aplicar_negrito_prefixo, runs_completos

_CATEGORIAS_DE_TITULO = {"titulo1", "titulo2", "titulo3", "titulo4", "titulo5", "titulo_sem_numero"}

_CORPO_SEM_RECUO = replace(config.CORPO_TEXTO, recuo_primeira_linha_cm=0.0)

# categoria -> (estilo, forcar_negrito_italico, forcar_maiusculas, estilo_word)
_REGRAS = {
    "titulo1": (config.TITULO_SECAO_PRIMARIA, True, True, "Heading 1"),
    "titulo2": (config.TITULO_SECAO_SECUNDARIA, True, True, "Heading 2"),
    "titulo3": (config.TITULO_SECAO_TERCIARIA, True, False, "Heading 3"),
    "titulo4": (config.TITULO_SECAO_QUATERNARIA, True, False, "Heading 4"),
    "titulo5": (config.TITULO_SECAO_QUINARIA, True, False, "Heading 5"),
    "titulo_sem_numero": (config.TITULO_SEM_NUMERO, True, True, "Título de Seção"),
    "legenda": (config.LEGENDA, True, False, None),
    "fonte_ilustracao": (config.FONTE_ILUSTRACAO, False, False, None),
    "citacao_longa": (config.CITACAO_LONGA, False, False, None),
    "referencia": (config.REFERENCIA, False, False, None),
    "resumo_corpo": (config.CORPO_TEXTO, False, False, None),
    "corpo": (config.CORPO_TEXTO, False, False, None),
    "corpo_sem_recuo": (_CORPO_SEM_RECUO, False, False, None),
}

# Apêndice I: "As seções primárias devem iniciar SEMPRE em páginas distintas."
# O Modelo de TCC oficial confirma que o mesmo vale para os elementos
# pré-textuais sem indicativo numérico (Dedicatória, Agradecimentos, Resumo,
# Abstract, Listas, Sumário...): no arquivo do modelo, cada um deles começa
# numa seção/página própria -- por isso "titulo_sem_numero" entra aqui
# também, não só "titulo1".
_CATEGORIAS_COM_QUEBRA_DE_PAGINA = {"titulo1", "titulo_sem_numero"}


@dataclass
class EventoFormatacao:
    indice: int
    categoria: str
    trecho: str


def _normalizar_apenas_fonte(paragraph, tamanho_pt: float = config.CORPO_TEXTO.tamanho_pt) -> None:
    for run in runs_completos(paragraph):
        run.font.name = config.CORPO_TEXTO.fonte
        run.font.size = Pt(tamanho_pt)


def _paragrafos_de_tabelas(tables):
    """Percorre todos os parágrafos dentro de células de tabela,
    recursivamente (uma célula pode conter outra tabela dentro)."""
    for tabela in tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                yield from celula.paragraphs
                yield from _paragrafos_de_tabelas(celula.tables)


def _formatar_tabelas(document) -> list[EventoFormatacao]:
    """Normaliza a fonte do conteúdo interno das tabelas (Apêndice I:
    tamanho 10). Não mexe em alinhamento/negrito -- só o tamanho é regra
    explícita para o miolo da tabela."""
    eventos: list[EventoFormatacao] = []
    for paragraph in _paragrafos_de_tabelas(document.tables):
        if not paragraph.text.strip():
            continue
        _normalizar_apenas_fonte(paragraph, config.TABELA_CONTEUDO_TAMANHO_PT)
        eventos.append(EventoFormatacao(indice=-1, categoria="tabela_conteudo", trecho=paragraph.text.strip()[:60]))
    return eventos


def _numId_ilvl(paragraph) -> Optional[tuple[str, int]]:
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        return None
    ilvl_el = numPr.find(qn("w:ilvl"))
    ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    return numId_el.get(qn("w:val")), ilvl


def _normalizar_fonte_nivel_numeracao(document, numId: str, ilvl: int, estilo, negrito: Optional[bool]) -> None:
    """Quando o título é numerado pela lista multinível automática do Word
    (ver classify.py), o número "1"/"1.1" que o Word desenha não usa a
    formatação do texto do parágrafo -- usa a formatação de fonte definida
    no próprio NÍVEL da lista (`abstractNum/lvl/rPr`, dentro de
    numbering.xml). Sem normalizar isso, o número sai na fonte que o nível
    já tinha antes (herdada de quando o aluno criou a lista), diferente da
    fonte do resto do título -- e se essa fonte antiga for, por exemplo,
    branca ou de tamanho quase zero, o número simplesmente não aparece."""
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return
    num_el = numbering.find(f'{qn("w:num")}[@{qn("w:numId")}="{numId}"]')
    if num_el is None:
        return
    abstract_ref = num_el.find(qn("w:abstractNumId"))
    if abstract_ref is None:
        return
    abstract_id = abstract_ref.get(qn("w:val"))
    abstract_el = numbering.find(f'{qn("w:abstractNum")}[@{qn("w:abstractNumId")}="{abstract_id}"]')
    if abstract_el is None:
        return
    lvl = abstract_el.find(f'{qn("w:lvl")}[@{qn("w:ilvl")}="{ilvl}"]')
    if lvl is None:
        return

    rPr = lvl.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        lvl.append(rPr)
    for filho in list(rPr):
        rPr.remove(filho)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), estilo.fonte)
    rFonts.set(qn("w:hAnsi"), estilo.fonte)
    rFonts.set(qn("w:cs"), estilo.fonte)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(estilo.tamanho_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(estilo.tamanho_pt * 2)))
    rPr.append(szCs)
    if negrito is not None:
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "1" if negrito else "0")
        rPr.append(b)
        bCs = OxmlElement("w:bCs")
        bCs.set(qn("w:val"), "1" if negrito else "0")
        rPr.append(bCs)


def _remover_quebra_de_pagina_manual(paragraph) -> bool:
    """Remove quebras de página manuais (Ctrl+Enter) de um parágrafo.

    A ferramenta já garante a quebra de página no lugar certo -- antes de
    cada seção primária (ver `_CATEGORIAS_COM_QUEBRA_DE_PAGINA`) -- por meio
    da propriedade de parágrafo `page_break_before`. Uma quebra manual
    (`<w:br w:type="page"/>`) solta no meio do texto é sempre redundante na
    melhor das hipóteses e está fora do lugar na pior, então é removida
    mesmo que o aluno a tenha colocado lá -- isso não altera nenhum texto,
    só o controle de quebra de linha/página.
    """
    removida = False
    for run in list(paragraph.runs):
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                run._element.remove(br)
                removida = True
    return removida


def _rebaixar_estilo_titulo_indevido(paragraph) -> None:
    """Se o parágrafo carrega um estilo Heading 1-5 mas a classificação
    concluiu que ele não é, de fato, um título (ex.: uma legenda ou um
    parágrafo de corpo que herdou o estilo por engano), tira esse estilo --
    senão o Sumário nativo do Word continuaria listando esse parágrafo como
    se fosse um capítulo, porque o campo TOC lê o nível de tópico do
    estilo, não a formatação visual aplicada por cima."""
    try:
        nome_atual = paragraph.style.name if paragraph.style else None
    except Exception:
        nome_atual = None
    if not eh_nome_de_estilo_titulo_numerado(nome_atual):
        return
    try:
        paragraph.style = paragraph.part.document.styles["Normal"]
    except KeyError:
        pass


def _tentar_aplicar_estilo_word(paragraph, nome_estilo: Optional[str]) -> None:
    if not nome_estilo:
        return
    try:
        paragraph.style = paragraph.part.document.styles[nome_estilo]
    except KeyError:
        # o documento do aluno não tem esse estilo definido -- a formatação
        # direta aplicada em seguida garante a aparência correta mesmo assim.
        pass


def formatar_documento(document) -> list[EventoFormatacao]:
    aplicar_configuracao_pagina(document, config.PAGE_SETUP)

    estado = EstadoClassificacao()
    eventos: list[EventoFormatacao] = []
    niveis_numeracao_normalizados: set[tuple[str, int]] = set()
    # Parágrafos em branco vistos "agora há pouco" (ver limpeza logo abaixo).
    buffer_paragrafos_vazios: list = []

    for indice, paragraph in enumerate(document.paragraphs):
        categoria = classificar_paragrafo(paragraph, estado)

        # Só preserva quebra de página manual em parágrafos "outro"/"vazio"
        # (capa, folha de rosto, ficha catalográfica, folha de aprovação...)
        # QUANDO ainda estamos na zona pré-textual (antes de qualquer
        # Dedicatória/Resumo/Sumário) -- ali a ferramenta não tem nenhum
        # jeito de saber onde uma página termina e a próxima começa, então
        # remover a quebra que o aluno colocou lá deixaria elementos
        # pré-textuais inteiros grudados, sem o aluno ter como corrigir
        # isso nem reprocessando o arquivo. Fora dessa zona (corpo do
        # trabalho), qualquer quebra manual continua sendo redundante ou
        # errada, porque `_CATEGORIAS_COM_QUEBRA_DE_PAGINA` cuida disso
        # sozinha -- mantém o comportamento original ali.
        preservar_quebra_manual = categoria in ("outro", "vazio") and em_zona_pre_textual(estado)
        if not preservar_quebra_manual and _remover_quebra_de_pagina_manual(paragraph):
            eventos.append(EventoFormatacao(indice=indice, categoria="quebra_pagina_removida", trecho=paragraph.text.strip()[:60]))

        if categoria == "vazio":
            buffer_paragrafos_vazios.append(paragraph._p)
            continue

        # Documentos reais contam só com parágrafos em branco empilhados
        # para empurrar Resumo/Abstract/Sumário/um capítulo para a próxima
        # página (sem quebra de página/seção de verdade). Agora que a
        # ferramenta insere sua própria quebra automática antes desses
        # títulos (`_CATEGORIAS_COM_QUEBRA_DE_PAGINA`), esses parágrafos em
        # branco viram um problema nunca antes: se sobrarem no fim da
        # página anterior e também "estourarem" para a página seguinte por
        # conta própria, o resultado é uma página inteira em branco entre
        # os dois elementos. Como não têm nenhum texto, removê-los aqui não
        # fere a regra de nunca alterar o conteúdo do trabalho.
        if categoria in _CATEGORIAS_COM_QUEBRA_DE_PAGINA and buffer_paragrafos_vazios:
            for elemento_vazio in buffer_paragrafos_vazios:
                elemento_vazio.getparent().remove(elemento_vazio)
            eventos.append(EventoFormatacao(
                indice=indice,
                categoria="paragrafos_vazios_removidos_antes_de_titulo",
                trecho=paragraph.text.strip()[:60],
            ))
        buffer_paragrafos_vazios = []

        if categoria == "sumario_entrada":
            continue

        if categoria not in _CATEGORIAS_DE_TITULO:
            _rebaixar_estilo_titulo_indevido(paragraph)

        if categoria == "outro":
            _normalizar_apenas_fonte(paragraph)
            if eh_marcador_ficha_catalografica(paragraph.text):
                paragraph.paragraph_format.page_break_before = True
            continue

        regra = _REGRAS.get(categoria)
        if regra is None:
            _normalizar_apenas_fonte(paragraph)
            continue

        estilo, forcar_negrito_italico, forcar_maiusculas, nome_estilo_word = regra
        _tentar_aplicar_estilo_word(paragraph, nome_estilo_word)
        aplicar_estilo(
            paragraph,
            estilo,
            forcar_fonte=True,
            forcar_negrito_italico=forcar_negrito_italico,
            forcar_maiusculas=forcar_maiusculas,
        )
        paragraph.paragraph_format.page_break_before = categoria in _CATEGORIAS_COM_QUEBRA_DE_PAGINA

        if categoria in ("titulo1", "titulo2", "titulo3", "titulo4", "titulo5"):
            chave = _numId_ilvl(paragraph)
            if chave is not None and chave not in niveis_numeracao_normalizados:
                niveis_numeracao_normalizados.add(chave)
                numId, ilvl = chave
                negrito = estilo.negrito if forcar_negrito_italico else None
                _normalizar_fonte_nivel_numeracao(document, numId, ilvl, estilo, negrito)

        if categoria == "fonte_ilustracao":
            aplicar_negrito_prefixo(paragraph, config.PREFIXO_FONTE_NEGRITO)

        trecho = paragraph.text.strip()[:60]
        eventos.append(EventoFormatacao(indice=indice, categoria=categoria, trecho=trecho))

    eventos.extend(_formatar_tabelas(document))

    return eventos
