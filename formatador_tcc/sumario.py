"""Reconstrói o SUMÁRIO como um campo TOC nativo do Word.

Em vez de tentar calcular números de página manualmente (o que exigiria um
motor de paginação real, algo que o python-docx não tem), inserimos o mesmo
campo `{ TOC }` que o Word usa nativamente: ele é recalculado pelo próprio
Word ao abrir o documento (via `updateFields`) ou com F9, sempre com os
números de página corretos, a partir dos títulos que o formatador marcou
como "Heading 1/2/3" (ver formatador.py). As entradas antigas do sumário
(texto colado ou desatualizado) são removidas.
"""
from __future__ import annotations

from dataclasses import dataclass

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import config
from .classify import EstadoClassificacao, classificar_paragrafo, MARCADOR_CAMPO_SUMARIO
from .texto import normalizar


@dataclass
class ResultadoSumario:
    encontrado: bool
    entradas_removidas: int = 0
    campo_inserido: bool = False


def _localizar_bloco_sumario(document):
    estado = EstadoClassificacao()
    indice_titulo = None
    indices_entradas: list[int] = []
    for i, paragraph in enumerate(document.paragraphs):
        categoria = classificar_paragrafo(paragraph, estado)
        if categoria == "titulo_sem_numero" and normalizar(paragraph.text) == "SUMARIO":
            indice_titulo = i
            continue
        if indice_titulo is not None and categoria == "sumario_entrada":
            indices_entradas.append(i)
        elif indice_titulo is not None and categoria != "sumario_entrada" and indices_entradas:
            break
    return indice_titulo, indices_entradas


# Blocos de conteúdo automático do Word (inseridos via Inserir -> Sumário /
# Índice de Ilustrações) que ficam "invisíveis" para document.paragraphs,
# porque moram dentro de um <w:sdt> (bloco de conteúdo), não como <w:p>
# soltos no corpo do documento. Se não forem removidos junto com as
# entradas antigas em texto puro, sobra um sumário nativo do Word inteiro
# (com seus próprios hyperlinks e marcadores internos) logo abaixo do novo
# campo TOC -- na prática isso já produziu um .docx que o Word se recusa a
# abrir ("problema com o conteúdo").
_GALERIAS_SUMARIO_NATIVO = {"Table of Contents", "Table of Figures"}


def _eh_sdt_de_sumario_nativo(elemento) -> bool:
    if elemento.tag != qn("w:sdt"):
        return False
    galeria = elemento.find(f"{qn('w:sdtPr')}/{qn('w:docPartObj')}/{qn('w:docPartGallery')}")
    if galeria is None:
        return False
    return galeria.get(qn("w:val")) in _GALERIAS_SUMARIO_NATIVO


def _remover_sumario_nativo_do_word(document, indice_titulo: int, indices_entradas: list[int]) -> int:
    """Remove qualquer sumário/índice nativo do Word (bloco de conteúdo
    "Sumário" ou "Índice de Ilustrações" inserido via Inserir -> Sumário)
    que esteja logo após o título SUMÁRIO -- percorre os irmãos no XML
    porque esses blocos não aparecem em document.paragraphs."""
    paragrafos = document.paragraphs
    titulo_p = paragrafos[indice_titulo]._p
    elementos_de_entradas = {paragrafos[i]._p for i in indices_entradas}

    removidos = 0
    for irmao in list(titulo_p.itersiblings()):
        if _eh_sdt_de_sumario_nativo(irmao):
            irmao.getparent().remove(irmao)
            removidos += 1
            continue
        if irmao.tag == qn("w:p") and irmao in elementos_de_entradas:
            continue
        break
    return removidos


def _forcar_atualizacao_de_campos_ao_abrir(document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


# Quando o Word recalcula o campo `{ TOC }`, cada entrada gerada usa o
# estilo de parágrafo embutido "TOC 1".."TOC 9" -- se o documento não os
# define, o Word cria versões padrão baseadas em "Normal" na hora, e como
# muitos modelos (inclusive o Modelo de TCC oficial) não sobrescrevem a
# fonte do "Normal", essas entradas saem na fonte/tamanho padrão do Word
# (ex.: Arial 11), não na fonte definida para o resto do trabalho. O
# Apêndice I é explícito: "Deve-se utilizar apenas um dos tipos [de fonte]
# escolhidos em todo o trabalho" (tamanho 12) -- por isso os estilos do
# Sumário são definidos aqui explicitamente com a mesma fonte do corpo do
# texto (config.CORPO_TEXTO), para as entradas geradas ficarem consistentes
# com o resto do documento em vez de usar a fonte "antiga" do template.
_NIVEIS_ESTILO_TOC = 5  # Apêndice I limita a numeração até a seção quinária.


def _garantir_estilos_sumario(document) -> None:
    styles_element = document.styles.element
    for nivel in range(1, _NIVEIS_ESTILO_TOC + 1):
        style_id = f"TOC{nivel}"
        estilo_existente = styles_element.find(f'{qn("w:style")}[@{qn("w:styleId")}="{style_id}"]')
        if estilo_existente is not None:
            rPr = estilo_existente.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                estilo_existente.append(rPr)
        else:
            estilo_existente = OxmlElement("w:style")
            estilo_existente.set(qn("w:type"), "paragraph")
            estilo_existente.set(qn("w:styleId"), style_id)
            nome = OxmlElement("w:name")
            nome.set(qn("w:val"), f"toc {nivel}")
            estilo_existente.append(nome)
            based_on = OxmlElement("w:basedOn")
            based_on.set(qn("w:val"), "Normal")
            estilo_existente.append(based_on)
            rPr = OxmlElement("w:rPr")
            estilo_existente.append(rPr)
            styles_element.append(estilo_existente)

        for filho in list(rPr):
            rPr.remove(filho)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), config.CORPO_TEXTO.fonte)
        rFonts.set(qn("w:hAnsi"), config.CORPO_TEXTO.fonte)
        rFonts.set(qn("w:cs"), config.CORPO_TEXTO.fonte)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(config.CORPO_TEXTO.tamanho_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(config.CORPO_TEXTO.tamanho_pt * 2)))
        rPr.append(szCs)


def _construir_campo_toc(paragraph) -> None:
    """Substitui o conteúdo do parágrafo por um campo `{ TOC \\o "1-5" \\h \\z \\u }`."""
    p = paragraph._p
    for child in list(p.findall(qn("w:r"))):
        p.remove(child)
    for child in list(p.findall(qn("w:hyperlink"))):
        p.remove(child)

    def novo_run():
        r = OxmlElement("w:r")
        p.append(r)
        return r

    r1 = novo_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    r1.append(begin)

    r2 = novo_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-5" \\h \\z \\u '
    r2.append(instr)

    r3 = novo_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3.append(sep)

    r4 = novo_run()
    t = OxmlElement("w:t")
    t.text = f"{MARCADOR_CAMPO_SUMARIO} -- clique com o botão direito e escolha “Atualizar campo” (ou F9) após abrir no Word."
    r4.append(t)

    r5 = novo_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5.append(end)


def reconstruir_sumario(document) -> ResultadoSumario:
    indice_titulo, indices_entradas = _localizar_bloco_sumario(document)
    if indice_titulo is None:
        return ResultadoSumario(encontrado=False)

    # precisa rodar ANTES de inserir o novo campo -- percorre os irmãos do
    # título no XML original, e o parágrafo novo entraria no meio do
    # caminho se já tivesse sido inserido.
    _remover_sumario_nativo_do_word(document, indice_titulo, indices_entradas)

    paragrafos = document.paragraphs
    titulo = paragrafos[indice_titulo]

    novo_p = OxmlElement("w:p")
    titulo._p.addnext(novo_p)
    from docx.text.paragraph import Paragraph
    campo_paragraph = Paragraph(novo_p, titulo._parent)
    _construir_campo_toc(campo_paragraph)

    removidos = 0
    for i in indices_entradas:
        alvo = paragrafos[i]._p
        alvo.getparent().remove(alvo)
        removidos += 1

    _garantir_estilos_sumario(document)
    _forcar_atualizacao_de_campos_ao_abrir(document)

    return ResultadoSumario(encontrado=True, entradas_removidas=removidos, campo_inserido=True)
