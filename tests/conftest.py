import docx
import pytest


def _heading(document, texto, nivel=1):
    p = document.add_paragraph(texto)
    p.style = document.styles[f"Heading {nivel}"]
    return p


@pytest.fixture
def construir_docx():
    """Fábrica de documentos .docx sintéticos para os testes."""

    def _construir(estrutura: list[tuple[str, str]]):
        """estrutura: lista de (tipo, texto).

        tipo pode ser: 'titulo_sem_numero', 'heading1', 'heading2', 'heading3',
        'texto' (parágrafo comum) ou 'referencia' (alias de 'texto', só documenta a intenção).
        """
        d = docx.Document()
        for tipo, texto in estrutura:
            if tipo in ("heading1", "titulo1"):
                _heading(d, texto, 1)
            elif tipo in ("heading2", "titulo2"):
                _heading(d, texto, 2)
            elif tipo in ("heading3", "titulo3"):
                _heading(d, texto, 3)
            elif tipo in ("heading4", "titulo4"):
                _heading(d, texto, 4)
            elif tipo in ("heading5", "titulo5"):
                _heading(d, texto, 5)
            else:
                d.add_paragraph(texto)
        return d

    return _construir
