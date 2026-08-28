from pathlib import Path

import docx
from docx.shared import Cm

from formatador_tcc.classify import EstadoClassificacao, classificar_paragrafo
from formatador_tcc.formatador import formatar_documento
from formatador_tcc.sumario import reconstruir_sumario

TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "Modelo_TCC_oficial.docx"


def test_formatar_nao_altera_texto(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Um parágrafo qualquer com texto do aluno que não deve mudar."),
        ("heading1", "REFERÊNCIAS"),
        ("texto", "SILVA, João. Um livro. São Paulo: Editora, 2020."),
    ])
    textos_antes = [p.text for p in d.paragraphs]
    formatar_documento(d)
    textos_depois = [p.text for p in d.paragraphs]
    assert textos_antes == textos_depois


def test_titulo_numerado_digitado_manualmente_e_reconhecido(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("texto", "1 INTRODUÇÃO"),
        ("texto", "Texto do primeiro parágrafo."),
        ("texto", "3.1 Um subtítulo qualquer"),
    ])
    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]
    assert categorias[1] == "titulo1"
    assert categorias[3] == "titulo2"


def test_paragrafo_corpo_fica_justificado_com_recuo(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Primeiro parágrafo logo após o título."),
        ("texto", "Segundo parágrafo, este sim com recuo de primeira linha."),
    ])
    formatar_documento(d)
    paragrafos = d.paragraphs
    # segundo parágrafo de corpo (índice 3) deve ter recuo ~1.25cm
    recuo = paragrafos[3].paragraph_format.first_line_indent
    assert recuo is not None
    assert abs(recuo.cm - 1.25) < 0.05


def test_titulo_ate_quinto_nivel_reconhecido(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("texto", "1 INTRODUÇÃO"),
        ("texto", "1.1 Contextualização"),
        ("texto", "1.1.1 Um Subtema"),
        ("texto", "1.1.1.1 Detalhe do subtema"),
        ("texto", "1.1.1.1.1 Último nível permitido"),
    ])
    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]
    assert categorias[1:6] == ["titulo1", "titulo2", "titulo3", "titulo4", "titulo5"]


def test_secao_secundaria_fica_caixa_alta_sem_negrito(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("heading2", "Contextualização"),
    ])
    formatar_documento(d)
    p = d.paragraphs[2]
    for run in p.runs:
        assert run.font.bold is False
        assert run.font.all_caps is True


def test_quebra_de_pagina_so_em_titulo1(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("heading2", "Contextualização"),
        ("texto", "Um parágrafo de corpo qualquer."),
    ])
    formatar_documento(d)
    assert d.paragraphs[1].paragraph_format.page_break_before is True
    assert d.paragraphs[2].paragraph_format.page_break_before is False
    assert d.paragraphs[3].paragraph_format.page_break_before is False


def test_quebra_de_pagina_tambem_em_titulo_sem_numero(construir_docx):
    """Regressão: um TCC real relia só em parágrafos em branco (não em
    quebra de página/seção real) para separar Dedicatória, Agradecimentos,
    Resumo, Abstract, Listas e Sumário -- ao reformatar (fonte/espaçamento
    diferentes), esses elementos deixavam de começar em página própria. O
    Modelo de TCC oficial mostra que cada um desses elementos pré-textuais
    começa em página/seção própria, igual às seções primárias."""
    d = construir_docx([
        ("texto", "Capa (texto livre, não classificável)."),
        ("titulo_sem_numero", "RESUMO"),
        ("texto", "Texto do resumo."),
        ("titulo_sem_numero", "SUMÁRIO"),
    ])
    formatar_documento(d)
    assert d.paragraphs[1].paragraph_format.page_break_before is True
    assert d.paragraphs[2].paragraph_format.page_break_before is False
    assert d.paragraphs[3].paragraph_format.page_break_before is True


def test_paragrafos_vazios_antes_de_titulo_sao_removidos(construir_docx):
    """Regressão real: um TCC empilhava vários parágrafos em branco entre o
    fim do Resumo e o título ABSTRACT (truque antigo do aluno para empurrar
    o Abstract para a página seguinte, sem quebra de página de verdade).
    Como a ferramenta agora insere sua própria quebra automática antes de
    ABSTRACT, esses parágrafos em branco sobrando criavam uma página
    inteira em branco entre os dois. Removê-los não fere o texto porque
    eles não têm nenhum conteúdo."""
    d = construir_docx([
        ("titulo_sem_numero", "RESUMO"),
        ("texto", "Texto do resumo."),
        ("texto", ""),
        ("texto", ""),
        ("texto", ""),
        ("titulo_sem_numero", "ABSTRACT"),
        ("texto", "Abstract text."),
    ])
    formatar_documento(d)

    textos = [p.text for p in d.paragraphs]
    assert textos == ["RESUMO", "Texto do resumo.", "ABSTRACT", "Abstract text."]
    idx_abstract = textos.index("ABSTRACT")
    assert d.paragraphs[idx_abstract].paragraph_format.page_break_before is True


def test_paragrafos_vazios_fora_de_zona_pretextual_nao_sao_removidos(construir_docx):
    """Parágrafos em branco que não estão logo antes de um título com
    quebra automática (ex.: espaçamento normal no meio do corpo do texto)
    continuam intocados -- a limpeza só vale para o caso específico do
    "empurrão" para a página de um título."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Primeiro parágrafo."),
        ("texto", ""),
        ("texto", ""),
        ("texto", "Segundo parágrafo, depois de linhas em branco no meio do corpo."),
    ])
    formatar_documento(d)

    textos = [p.text for p in d.paragraphs]
    assert textos == [
        "SUMÁRIO", "INTRODUÇÃO", "Primeiro parágrafo.", "", "",
        "Segundo parágrafo, depois de linhas em branco no meio do corpo.",
    ]


def test_ficha_catalografica_ganha_quebra_de_pagina(construir_docx):
    """A folha da Ficha Catalográfica é um elemento pré-textual obrigatório
    que deve começar em página própria -- diferente do título do trabalho
    ou nome do autor (texto livre do aluno, sem padrão confiável), o texto
    dessa folha é sempre um aviso fixo mencionando "Ficha Catalográfica",
    então dá pra reconhecer com segurança."""
    d = construir_docx([
        ("texto", "Capa (texto livre)."),
        ("texto", "Folha destinada à inclusão da Ficha Catalográfica (elemento obrigatório)."),
    ])
    formatar_documento(d)
    assert d.paragraphs[1].paragraph_format.page_break_before is True


def test_quebra_manual_entre_capa_e_contracapa_e_preservada(construir_docx):
    """Regressão: capa/contracapa/ficha catalográfica não têm nenhum
    padrão de texto que a ferramenta possa reconhecer -- então, se o aluno
    já colocou uma quebra de página manual ali (a única forma de garantir
    que cada um comece em página própria), a ferramenta não deve apagá-la
    (isso deixaria elementos pré-textuais inteiros grudados, sem o aluno
    ter como corrigir nem reprocessando o arquivo)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = construir_docx([
        ("texto", "Capa (texto livre)."),
        ("texto", ""),
    ])
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    d.paragraphs[1].add_run()._element.append(br)

    formatar_documento(d)

    tem_quebra = any(
        b.get(qn("w:type")) == "page"
        for p in d.paragraphs for r in p.runs for b in r._element.findall(qn("w:br"))
    )
    assert tem_quebra is True


def test_fonte_ilustracao_negrito_apenas_no_prefixo(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Figura 1 - Um gráfico qualquer"),
        ("texto", "Fonte: Silva e Souza, 2020."),
    ])
    formatar_documento(d)
    p = d.paragraphs[3]
    assert p.text == "Fonte: Silva e Souza, 2020."
    primeiro_run = p.runs[0]
    assert primeiro_run.text == "Fonte:"
    assert primeiro_run.font.bold is True
    assert any(r.font.bold is not True for r in p.runs[1:])


def test_toc_cache_de_referencias_nao_e_confundido_com_titulo(construir_docx):
    """Regressão: entradas antigas do sumário como 'REFERÊNCIAS\\t22' não
    podem ser classificadas como titulo1 -- senão sobrevivem à reconstrução
    do sumário como um "REFERÊNCIAS" duplicado e mal formatado."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("texto", "1\tINTRODUÇÃO\t12"),
        ("texto", "REFERÊNCIAS\t22"),
        ("texto", "APÊNDICE A -\t24"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Corpo do texto."),
        ("heading1", "REFERÊNCIAS"),
        ("texto", "SILVA, João. Um livro. São Paulo: Editora, 2020."),
    ])
    formatar_documento(d)
    resultado = reconstruir_sumario(d)
    assert resultado.entradas_removidas == 3

    textos = [p.text for p in d.paragraphs]
    assert textos.count("REFERÊNCIAS") == 1


def test_paginas_pretextuais_em_algarismo_romano_nao_confundem_sumario(construir_docx):
    """Regressão: um TCC real de aluno tinha as páginas pré-textuais (antes
    da Introdução) numeradas em algarismos romanos (ex.: "Sumário\\tix"),
    convenção comum em ABNT. Como RE_LINHA_SUMARIO só reconhecia dígitos
    decimais, essas linhas "saíam" da zona de SUMÁRIO cedo demais e a
    entrada de cache "1\\tIntrodução\\t1" era confundida com o título real,
    fazendo a numeração de página (e a quebra de página) cair na seção
    errada, bem no início do documento."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("texto", "Sumário\tix"),
        ("texto", "Lista de Figuras\txii"),
        ("texto", "1\tIntrodução\t1"),
        ("texto", "2\tObjetivos\t2"),
        ("heading1", "Introdução"),
        ("texto", "Corpo do texto."),
        ("heading1", "Objetivos"),
        ("texto", "Corpo do texto."),
    ])
    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]
    # nenhuma das linhas de cache deve virar título
    assert categorias[1:5] == ["sumario_entrada"] * 4
    # só os Heading 1 reais (índices 5 e 7) são titulo1
    assert categorias[5] == "titulo1"
    assert categorias[7] == "titulo1"


def test_legenda_com_estilo_de_titulo_nao_vira_titulo(construir_docx):
    """Regressão: um TCC real tinha a legenda 'Tabela 2.1 Valores...' com o
    estilo 'Heading 5' aplicado por engano (colagem de outro documento). A
    ferramenta confiava cegamente no estilo e tratava como um título de
    seção quinária -- inclusive entrando no Sumário."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("heading5", "Tabela 2.1 Valores codificados e reais das variáveis"),
    ])
    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]
    assert categorias[2] == "legenda"


def test_paragrafo_de_corpo_com_estilo_de_titulo_nao_vira_titulo(construir_docx):
    """Regressão: um parágrafo de corpo inteiro (frase longa) tinha o
    estilo 'Heading 2' aplicado por engano -- a ferramenta confiava
    cegamente no estilo e formatava o parágrafo inteiro como título
    (inclusive em CAIXA ALTA), destruindo a legibilidade do texto."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("heading2", "Com a finalidade de estudar o impacto dos espessantes "
                      "goma xantana e CMC nas propriedades físico-químicas "
                      "do produto final, foram realizados diversos testes."),
    ])
    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]
    assert categorias[2] in ("corpo", "corpo_sem_recuo")


def test_quebra_de_pagina_manual_e_removida(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto antes da quebra."),
    ])
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragrafo_vazio = d.add_paragraph()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    paragrafo_vazio.add_run()._element.append(br)

    assert any(
        r._element.findall(qn("w:br")) for p in d.paragraphs for r in p.runs
    )

    formatar_documento(d)

    assert not any(
        b.get(qn("w:type")) == "page"
        for p in d.paragraphs for r in p.runs for b in r._element.findall(qn("w:br"))
    )


def test_tabela_tem_fonte_normalizada(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
    ])
    tabela = d.add_table(rows=1, cols=1)
    run = tabela.rows[0].cells[0].paragraphs[0].add_run("Conteúdo da célula")
    run.font.size = None

    formatar_documento(d)

    celula_run = tabela.rows[0].cells[0].paragraphs[0].runs[0]
    assert celula_run.font.name == "Times New Roman"
    assert celula_run.font.size.pt == 10.0


def test_fonte_e_normalizada_dentro_de_hyperlink(construir_docx):
    """Regressão real: uma Lista de Tabelas com referência cruzada
    automática (Inserir > Referência Cruzada, "inserir como hyperlink")
    tinha os runs dentro de <w:hyperlink> -- paragraph.runs do python-docx
    só enxerga <w:r> soltos direto no parágrafo, então esse texto nunca
    recebia nenhuma normalização de fonte, mesmo depois de "formatado"."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
    ])
    p = d.add_paragraph()
    hyperlink = OxmlElement("w:hyperlink")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Comic Sans MS")
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Capa (texto livre com link)"
    r.append(t)
    hyperlink.append(r)
    p._p.append(hyperlink)

    formatar_documento(d)

    run_no_hyperlink = p._p.findall(f".//{qn('w:hyperlink')}/{qn('w:r')}")[0]
    rFonts_final = run_no_hyperlink.find(f"{qn('w:rPr')}/{qn('w:rFonts')}")
    assert rFonts_final.get(qn("w:ascii")) == "Times New Roman"


def test_legenda_com_estilo_de_titulo_perde_o_estilo_heading(construir_docx):
    """Regressão: não basta classificar corretamente como "legenda" -- se o
    parágrafo continuar com o estilo Word 'Heading 5', o campo de Sumário
    nativo (que lê o nível de tópico do estilo, não a formatação visual)
    continuaria listando a legenda como se fosse um título de seção."""
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("heading5", "Tabela 2.1 Valores codificados e reais das variáveis"),
    ])
    formatar_documento(d)
    assert d.paragraphs[2].style.name == "Normal"


def test_sumario_reconstroi_campo_toc(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("texto", "1\tINTRODUÇÃO\t12"),
        ("texto", "2\tOBJETIVOS\t13"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Corpo do texto."),
    ])
    formatar_documento(d)
    resultado = reconstruir_sumario(d)
    assert resultado.encontrado is True
    assert resultado.entradas_removidas == 2
    textos = [p.text for p in d.paragraphs]
    assert "1\tINTRODUÇÃO\t12" not in textos
    assert any("Sumário gerado automaticamente" in t for t in textos)


def test_documento_oficial_end_to_end():
    if not TEMPLATE.exists():
        return  # template não versionado neste checkout -- pula silenciosamente
    d = docx.Document(str(TEMPLATE))

    estado = EstadoClassificacao()
    textos_esperados = [p.text for p in d.paragraphs if classificar_paragrafo(p, estado) != "sumario_entrada"]

    eventos = formatar_documento(d)
    assert len(eventos) > 50  # documento extenso, deve reconhecer bastante coisa

    resultado_sumario = reconstruir_sumario(d)
    assert resultado_sumario.encontrado is True

    textos_finais = [p.text for p in d.paragraphs if "Sumário gerado automaticamente" not in p.text]
    # parágrafos em branco redundantes logo antes de um título com quebra de
    # página automática podem ser removidos (ver formatador.py) -- não
    # carregam texto, então isso não fere a garantia de conteúdo preservado;
    # o que importa aqui é que nenhum parágrafo COM TEXTO foi criado,
    # removido ou reordenado.
    assert [t for t in textos_finais if t.strip()] == [t for t in textos_esperados if t.strip()]
