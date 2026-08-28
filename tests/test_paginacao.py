import docx
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement

from formatador_tcc.paginacao import aplicar_numeracao_paginas


def _construir_docx_com_secoes():
    """Documento com as quebras de seção já no lugar certo (como o modelo
    oficial), incluindo uma antes da Introdução."""
    d = docx.Document()
    d.add_paragraph("CAPA")
    d.add_section(WD_SECTION.NEW_PAGE)
    d.add_paragraph("FICHA CATALOGRÁFICA")
    d.add_section(WD_SECTION.NEW_PAGE)
    d.add_paragraph("SUMÁRIO")
    d.add_section(WD_SECTION.NEW_PAGE)
    p = d.add_paragraph("INTRODUÇÃO")
    p.style = d.styles["Heading 1"]
    d.add_paragraph("Texto do primeiro parágrafo.")
    p2 = d.add_paragraph("REFERÊNCIAS")
    p2.style = d.styles["Heading 1"]
    d.add_paragraph("SILVA, João. Um livro. São Paulo: Editora, 2020.")
    return d


def _construir_docx_secao_unica():
    """Documento inteiro numa seção só, sem nenhuma quebra de seção --
    situação real de TCCs de alunos que não preservaram as quebras do
    modelo oficial ao digitar o trabalho."""
    d = docx.Document()
    d.add_paragraph("CAPA")
    d.add_paragraph("FICHA CATALOGRÁFICA")
    d.add_paragraph("SUMÁRIO")
    p = d.add_paragraph("INTRODUÇÃO")
    p.style = d.styles["Heading 1"]
    d.add_paragraph("Texto do primeiro parágrafo.")
    return d


def test_numeracao_aplicada_a_partir_da_introducao():
    d = _construir_docx_com_secoes()
    resultado = aplicar_numeracao_paginas(d)

    assert resultado.aplicada is True
    # a Introdução está na última seção (3 quebras de seção antes dela:
    # Capa | Ficha Catalográfica | Sumário | Introdução...)
    assert resultado.secao_introducao == 3

    secoes = d.sections
    # seções anteriores continuam sem número de página visível (o cabeçalho
    # da seção 0 pode reportar is_linked_to_previous de forma inconsistente
    # -- é uma particularidade da primeira seção do documento -- por isso a
    # checagem que importa de fato é o conteúdo do cabeçalho estar vazio)
    for s in secoes[:3]:
        assert "".join(p.text for p in s.header.paragraphs) == ""
    # a seção da Introdução tem cabeçalho próprio com o campo de página
    assert secoes[3].header.is_linked_to_previous is False


def test_campo_de_pagina_tem_formula_com_offset_correto():
    d = _construir_docx_com_secoes()
    aplicar_numeracao_paginas(d)

    header = d.sections[3].header
    xml = header._element.xml
    assert " PAGE " in xml and " - 2 " in xml and xml.count('fldCharType="begin"') == 2
    assert 'w:jc w:val="right"' in xml


def test_nao_duplica_campo_se_ja_houver_conteudo_no_cabecalho(construir_docx):
    d = _construir_docx_com_secoes()
    # simula um bloco de número de página já existente (como no modelo oficial)
    secao = d.sections[3]
    secao.header.is_linked_to_previous = False
    secao.header.add_paragraph("2")

    aplicar_numeracao_paginas(d)

    header = d.sections[3].header
    assert len(header.paragraphs) == 1
    assert " PAGE " in header._element.xml and " - 2 " in header._element.xml


def test_documento_de_secao_unica_ganha_quebra_de_secao_na_introducao():
    """Regressão: um TCC real de aluno tinha o documento inteiro numa única
    seção (nenhuma quebra de seção preservada do modelo oficial). Como
    cabeçalho é uma propriedade de seção, sem dividir em duas seções não
    havia como mostrar o número só a partir da Introdução -- ou aparecia
    também na capa, ou não aparecia em lugar nenhum."""
    d = _construir_docx_secao_unica()
    assert len(d.sections) == 1

    resultado = aplicar_numeracao_paginas(d)

    assert resultado.aplicada is True
    assert len(d.sections) == 2
    assert resultado.secao_introducao == 1

    secoes = d.sections
    assert "".join(p.text for p in secoes[0].header.paragraphs) == ""
    assert secoes[1].header.is_linked_to_previous is False
    assert " PAGE " in secoes[1].header._element.xml and " - 2 " in secoes[1].header._element.xml

    # o texto continua intacto, nenhum parágrafo foi criado/removido
    textos = [p.text for p in d.paragraphs]
    assert textos == ["CAPA", "FICHA CATALOGRÁFICA", "SUMÁRIO", "INTRODUÇÃO",
                       "Texto do primeiro parágrafo."]


def test_limpa_campo_de_pagina_deixado_em_secao_pretextual_por_execucao_anterior():
    """Regressão: um TCC real, já processado uma vez por uma versão com bug
    (que colocava o campo de página na seção errada, lá na Capa), foi
    reprocessado com a versão corrigida. O campo antigo, errado, continuava
    lá na seção pré-textual -- precisa ser limpo, não só ignorado."""
    d = _construir_docx_secao_unica()
    # simula o estado deixado por uma execução anterior com bug: campo de
    # página na ÚNICA seção existente (que cobre capa + tudo mais)
    secao = d.sections[0]
    secao.header.is_linked_to_previous = False
    secao.header.add_paragraph("2")

    resultado = aplicar_numeracao_paginas(d)

    assert resultado.aplicada is True
    secoes = d.sections
    # a seção pré-textual (agora seção 0, separada da Introdução) não deve
    # mais ter o "2" órfão da execução anterior
    assert "".join(p.text for p in secoes[0].header.paragraphs) == ""
    assert " PAGE " in secoes[1].header._element.xml and " - 2 " in secoes[1].header._element.xml


def test_campo_de_pagina_aninha_page_corretamente():
    """Regressão crítica: a primeira versão escrevia a instrução do campo
    de fórmula como texto simples " =PAGE-2 " (uma string só). O Word só
    reconhece PAGE como o campo nativo de número de página quando ele está
    ANINHADO como um campo de verdade (outro par begin/instrText/end)
    dentro da fórmula -- escrito como texto simples, "PAGE" é interpretado
    como o nome de um indicador (bookmark) inexistente, e o campo mostra
    "Erro! Indicador não definido." em vez do número (confirmado abrindo o
    arquivo gerado no Word de verdade)."""
    from docx.oxml.ns import qn

    d = _construir_docx_com_secoes()
    aplicar_numeracao_paginas(d)

    header = d.sections[3].header
    runs = header.paragraphs[0]._p.findall(f".//{qn('w:r')}")

    tipos_e_textos = []
    for r in runs:
        fld = r.find(qn("w:fldChar"))
        if fld is not None:
            tipos_e_textos.append(("fld", fld.get(qn("w:fldCharType"))))
            continue
        instr = r.find(qn("w:instrText"))
        if instr is not None:
            tipos_e_textos.append(("instr", instr.text))

    # campo externo (fórmula) -- begin, depois "=", depois o campo PAGE
    # aninhado (begin/instrText/end) ANTES do "- 2", depois separate/end.
    assert tipos_e_textos == [
        ("fld", "begin"),
        ("instr", " = "),
        ("fld", "begin"),
        ("instr", " PAGE "),
        ("fld", "end"),
        ("instr", " - 2 "),
        ("fld", "separate"),
        ("fld", "end"),
    ]


def test_remove_reinicio_de_numeracao_de_todas_as_secoes():
    """Regressão: um TCC real trazia `w:pgNumType w:start="1"` na (única)
    seção original do aluno. Como python-docx.add_section copia essa
    propriedade, tanto a seção pré-textual criada quanto a seção da
    Introdução acabavam reiniciando a contagem de página em 1 -- e o campo
    `=PAGE-2` da Introdução, em vez de mostrar a página física real do
    documento (ex.: 13), mostrava um número baixo ou até negativo, porque
    contava só a partir do início da própria seção, não do documento
    inteiro."""
    from docx.oxml.ns import qn

    d = _construir_docx_secao_unica()
    # como o Word grava por padrão: reinício de numeração em 1.
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    d.sections[0]._sectPr.append(pgNumType)

    aplicar_numeracao_paginas(d)

    for secao in d.sections:
        assert secao._sectPr.find(qn("w:pgNumType")) is None


def test_documento_sem_titulo1_nao_aplica_numeracao():
    d = docx.Document()
    d.add_paragraph("Texto qualquer sem nenhum capítulo numerado.")
    resultado = aplicar_numeracao_paginas(d)
    assert resultado.aplicada is False


def test_documento_final_continua_valido_apos_numeracao(tmp_path):
    import zipfile
    from lxml import etree

    d = _construir_docx_com_secoes()
    aplicar_numeracao_paginas(d)
    caminho = tmp_path / "saida.docx"
    d.save(caminho)

    with zipfile.ZipFile(caminho) as z:
        assert z.testzip() is None
        for name in z.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                etree.fromstring(z.read(name))
