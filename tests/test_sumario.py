"""Testes do sumario.py -- em especial as duas causas confirmadas de um
.docx real que o Word se recusava a abrir depois de formatado:

1. O sumário nativo do Word (Inserir -> Sumário, que vira um bloco
   <w:sdt> "invisível" para document.paragraphs) não era removido, e
   ficava duplicado ao lado do novo campo TOC.
2. Reprocessar um arquivo já formatado por esta ferramenta inseria um
   SEGUNDO campo TOC ao lado do primeiro, porque o parágrafo-placeholder
   ("Sumário gerado automaticamente...") não era reconhecido como uma
   entrada de sumário a ser substituída.
"""
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from formatador_tcc.sumario import reconstruir_sumario


def _adicionar_sdt_sumario_nativo(document, apos_paragraph, galeria="Table of Contents"):
    """Simula o bloco <w:sdt> que o Word insere via Inserir -> Sumário:
    um "Table of Contents" building block com parágrafos (hyperlinks,
    PAGEREF etc.) dentro de <w:sdtContent> -- invisível para
    document.paragraphs porque não é um <w:p> solto no corpo."""
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    docPartObj = OxmlElement("w:docPartObj")
    docPartGallery = OxmlElement("w:docPartGallery")
    docPartGallery.set(qn("w:val"), galeria)
    docPartObj.append(docPartGallery)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    p_interno = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1 Introdução\t12"
    r.append(t)
    p_interno.append(r)
    sdtContent.append(p_interno)
    sdt.append(sdtContent)

    apos_paragraph._p.addnext(sdt)
    return sdt


def _construir_docx_com_sumario_nativo(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Corpo do texto."),
    ])
    titulo = d.paragraphs[0]
    _adicionar_sdt_sumario_nativo(d, titulo)
    return d


def test_remove_sumario_nativo_do_word(construir_docx):
    d = _construir_docx_com_sumario_nativo(construir_docx)

    xml_antes = d.element.body.xml
    assert "Table of Contents" in xml_antes

    resultado = reconstruir_sumario(d)

    assert resultado.encontrado is True
    xml_depois = d.element.body.xml
    assert "Table of Contents" not in xml_depois
    assert xml_depois.count("Sumário gerado automaticamente") == 1


def test_reprocessar_documento_ja_formatado_nao_duplica_campo(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Corpo do texto."),
    ])

    reconstruir_sumario(d)
    total_paragrafos_1a_passada = len(d.paragraphs)
    campos_1a_passada = sum(1 for p in d.paragraphs if "Sumário gerado automaticamente" in p.text)
    assert campos_1a_passada == 1

    # reprocessa o MESMO documento já formatado, como um aluno reenviando
    # o arquivo que baixou -- não pode crescer nem duplicar o campo.
    resultado_2a = reconstruir_sumario(d)

    assert len(d.paragraphs) == total_paragrafos_1a_passada
    campos_2a_passada = sum(1 for p in d.paragraphs if "Sumário gerado automaticamente" in p.text)
    assert campos_2a_passada == 1


def test_reprocessar_tres_vezes_permanece_estavel(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Corpo do texto."),
    ])

    reconstruir_sumario(d)
    tamanho_apos_1 = len(d.paragraphs)
    reconstruir_sumario(d)
    tamanho_apos_2 = len(d.paragraphs)
    reconstruir_sumario(d)
    tamanho_apos_3 = len(d.paragraphs)

    assert tamanho_apos_1 == tamanho_apos_2 == tamanho_apos_3
