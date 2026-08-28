"""Testes de classify.py para títulos numerados via lista multinível
automática do Word (Formatar > Lista de Vários Níveis).

Cenário real que motivou isso: um TCC real usava esse recurso nativo do
Word para numerar os capítulos (em vez de digitar "1 INTRODUÇÃO" ou aplicar
o estilo Heading 1) -- o número "1"/"1.1" não aparece no texto do parágrafo
nem no nome do estilo, só no nível da lista (w:numPr). Sem essa detecção, a
ferramenta simplesmente não via TÍTULO NENHUM no corpo do trabalho (só
"REFERÊNCIAS", que é pego por outro caminho) -- nada de negrito, quebra de
página ou entrada no Sumário para nenhum capítulo.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from formatador_tcc.classify import EstadoClassificacao, classificar_paragrafo


def _abstract_num_multinivel(abstract_id: str) -> OxmlElement:
    """Lista multinível "de verdade": cada nível acumula os contadores dos
    níveis anteriores (1., 1.1., 1.1.1. ...), igual à numeração progressiva
    do Apêndice I."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    for ilvl in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "".join(f"%{i+1}." for i in range(ilvl + 1)))
        lvl.append(lvlText)
        abstract_num.append(lvl)
    return abstract_num


def _abstract_num_marcadores(abstract_id: str) -> OxmlElement:
    """Lista de marcadores comum (bullet) -- não tem relação com títulos."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    for ilvl in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "bullet")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "●")
        lvl.append(lvlText)
        abstract_num.append(lvl)
    return abstract_num


def _abstract_num_flat(abstract_id: str) -> OxmlElement:
    """Lista numerada "de um nível só": todos os níveis usam decimal mas
    NENHUM acumula os contadores do nível anterior (não é hierárquica) --
    ex.: um aluno enumerando "1) item", "2) item" no corpo do texto."""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    for ilvl in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "%1)")
        lvl.append(lvlText)
        abstract_num.append(lvl)
    return abstract_num


def _registrar_lista(document, num_id: str, abstract_id: str, abstract_num) -> None:
    numbering = document.part.numbering_part.element
    numbering.append(abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)


def _aplicar_numPr(paragraph, num_id: str, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(numId_el)
    pPr.append(numPr)


def test_titulo_via_lista_multinivel_e_reconhecido(construir_docx):
    d = construir_docx([
        ("texto", "INTRODUÇÃO"),
        ("texto", "Texto de corpo qualquer, bem mais longo que um título."),
        ("texto", "OBJETIVOS ESPECÍFICOS"),
    ])
    _registrar_lista(d, "112", "202", _abstract_num_multinivel("202"))
    _aplicar_numPr(d.paragraphs[0], "112", 0)
    _aplicar_numPr(d.paragraphs[2], "112", 1)

    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]

    assert categorias[0] == "titulo1"
    assert categorias[2] == "titulo2"


def test_lista_de_marcadores_nao_vira_titulo(construir_docx):
    d = construir_docx([
        ("texto", "INTRODUÇÃO"),
        ("texto", "Item de lista"),
    ])
    _registrar_lista(d, "111", "200", _abstract_num_marcadores("200"))
    _aplicar_numPr(d.paragraphs[1], "111", 0)

    estado = EstadoClassificacao()
    categorias = [classificar_paragrafo(p, estado) for p in d.paragraphs]

    assert categorias[1] != "titulo1"


def test_lista_numerada_de_nivel_unico_nao_vira_titulo(construir_docx):
    """Uma lista numerada comum (não hierárquica) não deve ser confundida
    com a numeração progressiva de capítulos -- mesmo usando numFmt=decimal,
    porque o nível 1 dela não acumula os contadores do nível 0."""
    d = construir_docx([
        ("texto", "Primeiro passo da receita"),
    ])
    _registrar_lista(d, "113", "201", _abstract_num_flat("201"))
    _aplicar_numPr(d.paragraphs[0], "113", 0)

    estado = EstadoClassificacao()
    categoria = classificar_paragrafo(d.paragraphs[0], estado)

    assert categoria != "titulo1"
