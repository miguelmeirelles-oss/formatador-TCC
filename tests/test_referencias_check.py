from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from formatador_tcc.referencias_check import verificar_referencias


def test_referencia_bem_formatada_nao_gera_problema(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto (SILVA, 2020)."),
        ("heading1", "REFERÊNCIAS"),
        ("texto", "SILVA, João. Um livro qualquer. São Paulo: Editora, 2020."),
    ])
    rel = verificar_referencias(d)
    # única pendência aceitável: falta de negrito/itálico no título (aviso,
    # não é erro estrutural) -- não deve haver problemas de autor/ano/pontuação
    mensagens = [p.mensagem for p in rel.problemas]
    assert not any("MAIÚSCULAS" in m for m in mensagens)
    assert not any("ano de publicação" in m for m in mensagens)
    assert not any("ponto final" in m for m in mensagens)


def test_detecta_autor_sem_maiusculas_ano_ausente_e_sem_ponto(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto (Silva, 2020)."),
        ("heading1", "REFERÊNCIAS"),
        ("texto", "souza, maria. obra sem padrão nenhum"),
    ])
    rel = verificar_referencias(d)
    mensagens = " | ".join(p.mensagem for p in rel.problemas)
    assert "MAIÚSCULAS" in mensagens
    assert "ano de publicação" in mensagens
    assert "ponto final" in mensagens


def test_detecta_entidade_com_ponto_como_valida(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto (BRASIL, 2002)."),
        ("heading1", "REFERÊNCIAS"),
        ("texto", "BRASIL. Lei nº 10.406, de 10 de janeiro de 2002. Institui o Código Civil."),
    ])
    rel = verificar_referencias(d)
    mensagens = [p.mensagem for p in rel.problemas]
    assert not any("MAIÚSCULAS" in m for m in mensagens)


def test_detecta_alinhamento_recuo_e_espacamento_incorretos(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto (SILVA, 2020)."),
        ("heading1", "REFERÊNCIAS"),
    ])
    ref = d.add_paragraph("SILVA, João. Um livro qualquer. São Paulo: Editora, 2020.")
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref.paragraph_format.first_line_indent = Cm(1.25)
    ref.paragraph_format.line_spacing = 1.5

    rel = verificar_referencias(d)
    corrigiveis = [p.mensagem for p in rel.problemas if p.corrigivel_automaticamente]
    assert any("alinhamento" in m for m in corrigiveis)
    assert any("recuo" in m for m in corrigiveis)
    assert any("espaçamento" in m for m in corrigiveis)


def test_destaque_inconsistente_entre_entradas(construir_docx):
    d = construir_docx([
        ("titulo_sem_numero", "SUMÁRIO"),
        ("heading1", "INTRODUÇÃO"),
        ("texto", "Texto (SILVA, 2020) e (OLIVEIRA, 2019)."),
        ("heading1", "REFERÊNCIAS"),
    ])
    r1 = d.add_paragraph()
    r1.add_run("SILVA, João. ").bold = True
    r1.add_run("Título em negrito.").bold = True
    r1.add_run(" São Paulo: Editora, 2020.")

    r2 = d.add_paragraph()
    r2.add_run("OLIVEIRA, Ana. ")
    r2.add_run("Título em itálico.").italic = True
    r2.add_run(" Curitiba: Editora, 2019.")

    rel = verificar_referencias(d)
    assert any("não é consistente" in p.mensagem for p in rel.problemas)
